"""
Jarvis Document Generator - Export conversations and knowledge to documents

Supports:
- Markdown (always available)
- Plain Text
- HTML (for email-ready content)
- PDF (optional, requires weasyprint)
- DOCX (optional, requires python-docx)
"""
import os
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field
from io import BytesIO

from .observability import get_logger, log_with_context

logger = get_logger("jarvis.documents")

BRAIN_ROOT = Path(os.environ.get("BRAIN_ROOT", "/brain"))
TEMPLATES_PATH = BRAIN_ROOT / "system" / "prompts" / "templates"
EXPORTS_PATH = BRAIN_ROOT / "exports"

# Ensure exports directory exists
EXPORTS_PATH.mkdir(parents=True, exist_ok=True)


# ============ Data Classes ============

@dataclass
class DocumentContent:
    """Content structure for document generation"""
    title: str
    sections: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    footer: str = ""


@dataclass
class GeneratedDocument:
    """Result of document generation"""
    filename: str
    format: str
    content: bytes
    path: Optional[str] = None
    size_bytes: int = 0


# ============ Template Loading ============

def _load_template(template_name: str) -> Optional[str]:
    """Load a Jinja2 template from the templates directory"""
    template_path = TEMPLATES_PATH / f"{template_name}.md.j2"
    if not template_path.exists():
        # Try without .md extension
        template_path = TEMPLATES_PATH / f"{template_name}.j2"

    if template_path.exists():
        return template_path.read_text(encoding="utf-8")
    return None


def _render_template(template_str: str, context: Dict[str, Any]) -> str:
    """Render a Jinja2 template with context"""
    try:
        from jinja2 import Template, Environment, BaseLoader
        env = Environment(loader=BaseLoader())
        # Add custom filters
        env.filters['date'] = lambda d, fmt='%d.%m.%Y': d.strftime(fmt) if d else ''
        env.filters['datetime'] = lambda d, fmt='%d.%m.%Y %H:%M': d.strftime(fmt) if d else ''
        template = env.from_string(template_str)
        return template.render(**context)
    except ImportError:
        # Fallback: simple string replacement
        result = template_str
        for key, value in context.items():
            result = result.replace(f"{{{{ {key} }}}}", str(value))
        return result


# ============ Format Converters ============

def _markdown_to_html(markdown_text: str) -> str:
    """Convert markdown to HTML"""
    try:
        import markdown
        return markdown.markdown(
            markdown_text,
            extensions=['tables', 'fenced_code', 'nl2br']
        )
    except ImportError:
        # Basic fallback
        html = markdown_text
        html = html.replace('\n\n', '</p><p>')
        html = html.replace('\n', '<br>')
        html = f"<p>{html}</p>"
        return html


def _html_to_pdf(html_content: str, title: str = "Document") -> bytes:
    """Convert HTML to PDF using weasyprint"""
    try:
        from weasyprint import HTML, CSS

        # Wrap in basic HTML structure with styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
            <style>
                body {{ font-family: 'Helvetica', 'Arial', sans-serif; margin: 2cm; line-height: 1.6; }}
                h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                h2 {{ color: #555; margin-top: 1.5em; }}
                h3 {{ color: #666; }}
                code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 1em; border-radius: 5px; overflow-x: auto; }}
                table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background: #f4f4f4; }}
                blockquote {{ border-left: 4px solid #ddd; margin: 1em 0; padding-left: 1em; color: #666; }}
                .footer {{ margin-top: 2em; padding-top: 1em; border-top: 1px solid #ddd; font-size: 0.9em; color: #888; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        pdf_bytes = BytesIO()
        HTML(string=full_html).write_pdf(pdf_bytes)
        return pdf_bytes.getvalue()
    except ImportError:
        raise ImportError("weasyprint not installed. Install with: pip install weasyprint")


def _markdown_to_docx(markdown_text: str, title: str = "Document") -> bytes:
    """Convert markdown to DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse markdown and add content
        lines = markdown_text.split('\n')
        in_code_block = False
        code_content = []

        for line in lines:
            # Code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    code_content = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_content.append(line)
                continue

            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Numbered lists
            elif line and line[0].isdigit() and '. ' in line[:4]:
                text = line.split('. ', 1)[1] if '. ' in line else line
                doc.add_paragraph(text, style='List Number')
            # Regular paragraphs
            elif line.strip():
                # Handle bold and italic
                doc.add_paragraph(line)
            # Empty lines create paragraph breaks
            elif not line.strip():
                doc.add_paragraph()

        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()
    except ImportError:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")


# ============ Document Types ============

def generate_email_draft(
    subject: str,
    to: str,
    body: str,
    tone: str = "professional",
    cc: str = "",
    context: str = ""
) -> str:
    """Generate an email draft in markdown format"""
    template = _load_template("email_draft")

    if template:
        return _render_template(template, {
            "subject": subject,
            "to": to,
            "cc": cc,
            "body": body,
            "tone": tone,
            "context": context,
            "date": datetime.now()
        })

    # Fallback template
    draft = f"""# Email Draft

**To:** {to}
{"**CC:** " + cc if cc else ""}
**Subject:** {subject}

---

{body}

---
*Tone: {tone}*
*Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}*
"""
    return draft


def generate_meeting_summary(
    title: str,
    date: datetime,
    attendees: List[str],
    agenda: List[str],
    decisions: List[str],
    action_items: List[Dict[str, str]],
    notes: str = ""
) -> str:
    """Generate a meeting summary in markdown format"""
    template = _load_template("meeting_summary")

    if template:
        return _render_template(template, {
            "title": title,
            "date": date,
            "attendees": attendees,
            "agenda": agenda,
            "decisions": decisions,
            "action_items": action_items,
            "notes": notes
        })

    # Fallback template
    attendees_str = ", ".join(attendees)
    agenda_str = "\n".join(f"- {item}" for item in agenda)
    decisions_str = "\n".join(f"- {d}" for d in decisions)
    actions_str = "\n".join(
        f"- [ ] {a['task']} (@{a.get('owner', 'TBD')}, due: {a.get('due', 'TBD')})"
        for a in action_items
    )

    return f"""# {title}

**Date:** {date.strftime('%d.%m.%Y %H:%M')}
**Attendees:** {attendees_str}

## Agenda
{agenda_str}

## Decisions
{decisions_str}

## Action Items
{actions_str}

## Notes
{notes if notes else "_No additional notes_"}

---
*Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}*
"""


def generate_linkedin_post(
    topic: str,
    hook: str,
    story: str,
    learning: str,
    cta: str,
    hashtags: List[str]
) -> str:
    """Generate a LinkedIn post in markdown format"""
    template = _load_template("linkedin_post")

    if template:
        return _render_template(template, {
            "topic": topic,
            "hook": hook,
            "story": story,
            "learning": learning,
            "cta": cta,
            "hashtags": hashtags
        })

    hashtags_str = " ".join(f"#{h}" for h in hashtags)

    return f"""# LinkedIn Post: {topic}

---

{hook}

{story}

{learning}

{cta}

{hashtags_str}

---
*Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}*
"""


def generate_progress_report(
    title: str,
    period: str,
    domain: str,
    achievements: List[str],
    challenges: List[str],
    next_steps: List[str],
    metrics: Dict[str, Any] = None
) -> str:
    """Generate a progress report in markdown format"""
    template = _load_template("progress_report")

    if template:
        return _render_template(template, {
            "title": title,
            "period": period,
            "domain": domain,
            "achievements": achievements,
            "challenges": challenges,
            "next_steps": next_steps,
            "metrics": metrics or {},
            "date": datetime.now()
        })

    achievements_str = "\n".join(f"- ✅ {a}" for a in achievements)
    challenges_str = "\n".join(f"- ⚠️ {c}" for c in challenges)
    next_steps_str = "\n".join(f"- [ ] {s}" for s in next_steps)

    metrics_section = ""
    if metrics:
        metrics_lines = [f"| {k} | {v} |" for k, v in metrics.items()]
        metrics_section = f"""
## Metrics
| Metric | Value |
|--------|-------|
{chr(10).join(metrics_lines)}
"""

    return f"""# {title}

**Period:** {period}
**Domain:** {domain}

## Achievements
{achievements_str}

## Challenges
{challenges_str}
{metrics_section}
## Next Steps
{next_steps_str}

---
*Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}*
"""


def generate_presentation_outline(
    title: str,
    audience: str,
    duration_minutes: int,
    core_message: str,
    sections: List[Dict[str, Any]]
) -> str:
    """Generate a presentation outline in markdown format"""
    template = _load_template("presentation")

    if template:
        return _render_template(template, {
            "title": title,
            "audience": audience,
            "duration": duration_minutes,
            "core_message": core_message,
            "sections": sections,
            "date": datetime.now()
        })

    sections_str = ""
    for i, section in enumerate(sections, 1):
        section_title = section.get("title", f"Section {i}")
        section_duration = section.get("duration", "")
        section_content = section.get("content", [])
        speaker_notes = section.get("notes", "")

        content_str = "\n".join(f"   - {c}" for c in section_content)
        notes_str = f"\n   > *Speaker notes: {speaker_notes}*" if speaker_notes else ""

        sections_str += f"""
### {i}. {section_title} ({section_duration})
{content_str}{notes_str}
"""

    return f"""# {title}

**Audience:** {audience}
**Duration:** {duration_minutes} minutes
**Core Message:** {core_message}

---

## Outline
{sections_str}

---
*Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}*
"""


# ============ Export Functions ============

def export_conversation(
    messages: List[Dict[str, str]],
    title: str = "Conversation Export",
    format: str = "md"
) -> GeneratedDocument:
    """Export a conversation to a document"""

    # Build markdown content
    lines = [f"# {title}", "", f"*Exported: {datetime.now().strftime('%d.%m.%Y %H:%M')}*", "", "---", ""]

    for msg in messages:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        timestamp = msg.get("timestamp", "")

        if role == "user":
            lines.append(f"**User** {timestamp}")
            lines.append(f"> {content}")
        else:
            lines.append(f"**Jarvis** {timestamp}")
            lines.append(content)
        lines.append("")

    markdown_content = "\n".join(lines)

    return _export_to_format(markdown_content, title, format)


def export_knowledge(
    knowledge_items: List[Dict[str, Any]],
    title: str = "Knowledge Export",
    format: str = "md"
) -> GeneratedDocument:
    """Export knowledge items to a document"""

    lines = [f"# {title}", "", f"*Exported: {datetime.now().strftime('%d.%m.%Y %H:%M')}*", "", "---", ""]

    for item in knowledge_items:
        item_type = item.get("type", "note")
        content = item.get("content", "")
        source = item.get("source", "")
        created = item.get("created_at", "")

        lines.append(f"## {item_type.title()}")
        if source:
            lines.append(f"*Source: {source}*")
        if created:
            lines.append(f"*Created: {created}*")
        lines.append("")
        lines.append(content)
        lines.append("")
        lines.append("---")
        lines.append("")

    markdown_content = "\n".join(lines)

    return _export_to_format(markdown_content, title, format)


def _export_to_format(
    markdown_content: str,
    title: str,
    format: str
) -> GeneratedDocument:
    """Convert markdown content to specified format and save"""

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
    safe_title = safe_title.replace(' ', '_')[:50]

    if format == "md" or format == "markdown":
        filename = f"{safe_title}_{timestamp}.md"
        content = markdown_content.encode('utf-8')

    elif format == "txt" or format == "text":
        filename = f"{safe_title}_{timestamp}.txt"
        # Strip markdown formatting for plain text
        plain_text = markdown_content
        plain_text = plain_text.replace('**', '')
        plain_text = plain_text.replace('*', '')
        plain_text = plain_text.replace('# ', '')
        plain_text = plain_text.replace('## ', '')
        plain_text = plain_text.replace('### ', '')
        plain_text = plain_text.replace('> ', '')
        content = plain_text.encode('utf-8')

    elif format == "html":
        filename = f"{safe_title}_{timestamp}.html"
        html_content = _markdown_to_html(markdown_content)
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{ font-family: system-ui, -apple-system, sans-serif; max-width: 800px; margin: 2em auto; padding: 0 1em; line-height: 1.6; }}
        h1 {{ color: #333; }}
        h2 {{ color: #555; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
        blockquote {{ border-left: 4px solid #ddd; margin: 1em 0; padding-left: 1em; color: #666; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        content = full_html.encode('utf-8')

    elif format == "pdf":
        filename = f"{safe_title}_{timestamp}.pdf"
        html_content = _markdown_to_html(markdown_content)
        content = _html_to_pdf(html_content, title)

    elif format == "docx":
        filename = f"{safe_title}_{timestamp}.docx"
        content = _markdown_to_docx(markdown_content, title)

    else:
        raise ValueError(f"Unsupported format: {format}. Supported: md, txt, html, pdf, docx")

    # Save to file
    file_path = EXPORTS_PATH / filename
    file_path.write_bytes(content)

    log_with_context(logger, "info", "Document exported",
                    filename=filename, format=format, size=len(content))

    return GeneratedDocument(
        filename=filename,
        format=format,
        content=content,
        path=str(file_path),
        size_bytes=len(content)
    )


# ============ Quick Generate Functions ============

def quick_generate(
    doc_type: str,
    content: Dict[str, Any],
    format: str = "md"
) -> GeneratedDocument:
    """Quick generate a document by type"""

    generators = {
        "email": lambda c: generate_email_draft(
            subject=c.get("subject", ""),
            to=c.get("to", ""),
            body=c.get("body", ""),
            tone=c.get("tone", "professional"),
            cc=c.get("cc", ""),
            context=c.get("context", "")
        ),
        "meeting": lambda c: generate_meeting_summary(
            title=c.get("title", "Meeting Summary"),
            date=c.get("date", datetime.now()),
            attendees=c.get("attendees", []),
            agenda=c.get("agenda", []),
            decisions=c.get("decisions", []),
            action_items=c.get("action_items", []),
            notes=c.get("notes", "")
        ),
        "linkedin": lambda c: generate_linkedin_post(
            topic=c.get("topic", ""),
            hook=c.get("hook", ""),
            story=c.get("story", ""),
            learning=c.get("learning", ""),
            cta=c.get("cta", ""),
            hashtags=c.get("hashtags", [])
        ),
        "progress": lambda c: generate_progress_report(
            title=c.get("title", "Progress Report"),
            period=c.get("period", ""),
            domain=c.get("domain", ""),
            achievements=c.get("achievements", []),
            challenges=c.get("challenges", []),
            next_steps=c.get("next_steps", []),
            metrics=c.get("metrics", {})
        ),
        "presentation": lambda c: generate_presentation_outline(
            title=c.get("title", "Presentation"),
            audience=c.get("audience", ""),
            duration_minutes=c.get("duration", 10),
            core_message=c.get("core_message", ""),
            sections=c.get("sections", [])
        ),
    }

    if doc_type not in generators:
        raise ValueError(f"Unknown document type: {doc_type}. Available: {', '.join(generators.keys())}")

    markdown_content = generators[doc_type](content)
    title = content.get("title", doc_type.title())

    return _export_to_format(markdown_content, title, format)


def list_exports(limit: int = 20) -> List[Dict[str, Any]]:
    """List recent exports"""
    exports = []

    for path in sorted(EXPORTS_PATH.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True)[:limit]:
        if path.is_file():
            stat = path.stat()
            exports.append({
                "filename": path.name,
                "path": str(path),
                "size_bytes": stat.st_size,
                "created_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "format": path.suffix[1:] if path.suffix else "unknown"
            })

    return exports


def get_export(filename: str) -> Optional[bytes]:
    """Get export file content by filename"""
    file_path = EXPORTS_PATH / filename
    if file_path.exists():
        return file_path.read_bytes()
    return None
