"""
Universal Document Parser Service

Parses various file formats for knowledge base indexing:
- Text: TXT, MD, CSV, JSON, LOG
- Documents: PDF, DOCX
- Images: JPEG, PNG, GIF, WEBP (via Claude Vision)

All parsers return a unified ParsedDocument structure.
"""

import os
import json
import csv
import io
import base64
from typing import Optional, Dict, Any, List, Tuple
from dataclasses import dataclass, field
from datetime import datetime

from ..observability import get_logger

logger = get_logger("jarvis.document_parser")


@dataclass
class ParsedDocument:
    """Unified structure for all parsed documents."""
    content: str  # Extracted text content
    title: str  # Document title (from content or filename)
    file_type: str  # Detected file type
    metadata: Dict[str, Any] = field(default_factory=dict)
    preview: str = ""  # Short preview for display
    success: bool = True
    error: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "title": self.title,
            "file_type": self.file_type,
            "metadata": self.metadata,
            "preview": self.preview,
            "success": self.success,
            "error": self.error
        }


# Supported file extensions by category
TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".xml", ".yaml", ".yml", ".html", ".htm"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".doc"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"}

# All supported extensions
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS | IMAGE_EXTENSIONS


def get_file_type(filename: str) -> str:
    """Determine file type from extension."""
    ext = os.path.splitext(filename.lower())[1]

    if ext in {".txt", ".log"}:
        return "text"
    elif ext == ".md":
        return "markdown"
    elif ext == ".csv":
        return "csv"
    elif ext == ".json":
        return "json"
    elif ext in {".xml", ".html", ".htm"}:
        return "markup"
    elif ext in {".yaml", ".yml"}:
        return "yaml"
    elif ext == ".pdf":
        return "pdf"
    elif ext in {".docx", ".doc"}:
        return "docx"
    elif ext in IMAGE_EXTENSIONS:
        return "image"
    else:
        return "unknown"


def is_supported(filename: str) -> bool:
    """Check if file type is supported."""
    ext = os.path.splitext(filename.lower())[1]
    return ext in SUPPORTED_EXTENSIONS


def get_supported_formats() -> Dict[str, List[str]]:
    """Return supported formats grouped by category."""
    return {
        "text": [".txt", ".md", ".log"],
        "data": [".csv", ".json", ".xml", ".yaml", ".yml"],
        "documents": [".pdf", ".docx"],
        "images": [".jpg", ".jpeg", ".png", ".gif", ".webp"]
    }


# ============================================================================
# Text Parsers
# ============================================================================

def parse_text(content: bytes, filename: str) -> ParsedDocument:
    """Parse plain text files (.txt, .log, .md)."""
    try:
        # Try UTF-8 first, then Latin-1
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        # Extract title from first line or filename
        lines = text.strip().split("\n")
        title = lines[0][:100].strip() if lines else filename

        # For markdown, extract # heading
        if filename.lower().endswith(".md"):
            for line in lines[:10]:
                if line.startswith("# "):
                    title = line[2:].strip()
                    break

        preview = text[:500] + "..." if len(text) > 500 else text

        return ParsedDocument(
            content=text,
            title=title,
            file_type=get_file_type(filename),
            metadata={
                "char_count": len(text),
                "line_count": len(lines),
                "encoding": "utf-8"
            },
            preview=preview
        )

    except Exception as e:
        logger.error(f"Error parsing text file {filename}: {e}")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="text",
            success=False,
            error=str(e)
        )


def parse_csv(content: bytes, filename: str) -> ParsedDocument:
    """Parse CSV files to structured text."""
    try:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return ParsedDocument(
                content="",
                title=filename,
                file_type="csv",
                success=False,
                error="Empty CSV file"
            )

        # Use first row as headers
        headers = rows[0]
        data_rows = rows[1:]

        # Convert to readable text format
        output_lines = []
        output_lines.append(f"CSV: {filename}")
        output_lines.append(f"Columns: {', '.join(headers)}")
        output_lines.append(f"Rows: {len(data_rows)}")
        output_lines.append("")

        # Add data as key-value pairs (limit to first 100 rows for KB)
        for i, row in enumerate(data_rows[:100]):
            row_text = []
            for j, value in enumerate(row):
                if j < len(headers) and value.strip():
                    row_text.append(f"{headers[j]}: {value}")
            if row_text:
                output_lines.append(f"Row {i+1}: " + " | ".join(row_text))

        if len(data_rows) > 100:
            output_lines.append(f"... and {len(data_rows) - 100} more rows")

        full_content = "\n".join(output_lines)

        return ParsedDocument(
            content=full_content,
            title=filename,
            file_type="csv",
            metadata={
                "columns": headers,
                "row_count": len(data_rows),
                "column_count": len(headers)
            },
            preview=full_content[:500]
        )

    except Exception as e:
        logger.error(f"Error parsing CSV {filename}: {e}")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="csv",
            success=False,
            error=str(e)
        )


def parse_json(content: bytes, filename: str) -> ParsedDocument:
    """Parse JSON files to readable text."""
    try:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        data = json.loads(text)

        # Pretty print for readability
        formatted = json.dumps(data, indent=2, ensure_ascii=False)

        # Extract title from data if possible
        title = filename
        if isinstance(data, dict):
            title = data.get("title") or data.get("name") or filename

        return ParsedDocument(
            content=formatted,
            title=title,
            file_type="json",
            metadata={
                "type": type(data).__name__,
                "keys": list(data.keys()) if isinstance(data, dict) else None,
                "length": len(data) if isinstance(data, (list, dict)) else None
            },
            preview=formatted[:500]
        )

    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in {filename}: {e}")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="json",
            success=False,
            error=f"Invalid JSON: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error parsing JSON {filename}: {e}")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="json",
            success=False,
            error=str(e)
        )


# ============================================================================
# Document Parsers (PDF, DOCX)
# ============================================================================

def parse_pdf(content: bytes, filename: str) -> ParsedDocument:
    """Parse PDF files using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))

        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            return ParsedDocument(
                content="",
                title=filename,
                file_type="pdf",
                success=False,
                error="PDF contains no extractable text (might be scanned/image-based)"
            )

        # Extract title from metadata or first line
        title = filename
        if reader.metadata:
            title = reader.metadata.get("/Title") or filename

        return ParsedDocument(
            content=full_text,
            title=title,
            file_type="pdf",
            metadata={
                "page_count": len(reader.pages),
                "author": reader.metadata.get("/Author") if reader.metadata else None,
                "created": str(reader.metadata.get("/CreationDate")) if reader.metadata else None
            },
            preview=full_text[:500]
        )

    except ImportError:
        logger.error("pypdf not installed")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="pdf",
            success=False,
            error="PDF parsing not available (pypdf not installed)"
        )
    except Exception as e:
        logger.error(f"Error parsing PDF {filename}: {e}")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="pdf",
            success=False,
            error=str(e)
        )


def parse_docx(content: bytes, filename: str) -> ParsedDocument:
    """Parse DOCX files using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Check for headings
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    text_parts.append(f"{'#' * int(level) if level.isdigit() else '##'} {para.text}")
                else:
                    text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
            if table_text:
                text_parts.append("\n[Table]\n" + "\n".join(table_text) + "\n")

        full_text = "\n\n".join(text_parts)

        # Extract title from first heading or filename
        title = filename
        for para in doc.paragraphs[:5]:
            if para.style and para.style.name.startswith("Heading"):
                title = para.text
                break
            elif para.text.strip() and len(para.text) < 100:
                title = para.text
                break

        # Get core properties
        metadata = {}
        try:
            props = doc.core_properties
            metadata = {
                "author": props.author,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
                "title": props.title
            }
            if props.title:
                title = props.title
        except Exception:
            pass

        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        return ParsedDocument(
            content=full_text,
            title=title,
            file_type="docx",
            metadata=metadata,
            preview=full_text[:500]
        )

    except ImportError:
        logger.error("python-docx not installed")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="docx",
            success=False,
            error="DOCX parsing not available (python-docx not installed)"
        )
    except Exception as e:
        logger.error(f"Error parsing DOCX {filename}: {e}")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="docx",
            success=False,
            error=str(e)
        )


# ============================================================================
# Image Parser (via Claude Vision)
# ============================================================================

def parse_image_with_vision(
    content: bytes,
    filename: str,
    extract_text: bool = True,
    describe: bool = True
) -> ParsedDocument:
    """
    Parse image using Claude Vision API.

    Args:
        content: Image bytes
        filename: Original filename
        extract_text: Whether to extract text (OCR)
        describe: Whether to describe the image content
    """
    try:
        import anthropic
        from ..config import config

        # Determine media type
        ext = os.path.splitext(filename.lower())[1]
        media_types = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".webp": "image/webp"
        }
        media_type = media_types.get(ext, "image/jpeg")

        # Encode image
        image_data = base64.b64encode(content).decode("utf-8")

        # Build prompt based on options
        prompt_parts = []
        if extract_text:
            prompt_parts.append("Extract all text visible in this image (OCR).")
        if describe:
            prompt_parts.append("Describe the main content and structure of this image.")
        prompt_parts.append("Format the output clearly with sections if applicable.")

        prompt = " ".join(prompt_parts)

        # Call Claude Vision
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": image_data
                            }
                        },
                        {
                            "type": "text",
                            "text": prompt
                        }
                    ]
                }
            ]
        )

        extracted_text = response.content[0].text

        return ParsedDocument(
            content=extracted_text,
            title=f"Image: {filename}",
            file_type="image",
            metadata={
                "media_type": media_type,
                "size_bytes": len(content),
                "extracted_via": "claude_vision"
            },
            preview=extracted_text[:500]
        )

    except ImportError:
        logger.error("anthropic not installed")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="image",
            success=False,
            error="Image parsing not available (anthropic not installed)"
        )
    except Exception as e:
        logger.error(f"Error parsing image {filename}: {e}")
        return ParsedDocument(
            content="",
            title=filename,
            file_type="image",
            success=False,
            error=str(e)
        )


# ============================================================================
# Main Parser Entry Point
# ============================================================================

def parse_document(
    content: bytes,
    filename: str,
    use_vision_for_images: bool = True
) -> ParsedDocument:
    """
    Parse any supported document format.

    Args:
        content: File content as bytes
        filename: Original filename (used for type detection)
        use_vision_for_images: Whether to use Claude Vision for images

    Returns:
        ParsedDocument with extracted content and metadata
    """
    ext = os.path.splitext(filename.lower())[1]

    if not is_supported(filename):
        return ParsedDocument(
            content="",
            title=filename,
            file_type="unknown",
            success=False,
            error=f"Unsupported file type: {ext}"
        )

    # Route to appropriate parser
    if ext in {".txt", ".md", ".log"}:
        return parse_text(content, filename)

    elif ext == ".csv":
        return parse_csv(content, filename)

    elif ext == ".json":
        return parse_json(content, filename)

    elif ext in {".xml", ".html", ".htm", ".yaml", ".yml"}:
        # Parse as text for now
        return parse_text(content, filename)

    elif ext == ".pdf":
        return parse_pdf(content, filename)

    elif ext in {".docx", ".doc"}:
        if ext == ".doc":
            return ParsedDocument(
                content="",
                title=filename,
                file_type="doc",
                success=False,
                error="Legacy .doc format not supported. Please convert to .docx"
            )
        return parse_docx(content, filename)

    elif ext in IMAGE_EXTENSIONS:
        if use_vision_for_images:
            return parse_image_with_vision(content, filename)
        else:
            return ParsedDocument(
                content="[Image file - vision parsing disabled]",
                title=filename,
                file_type="image",
                metadata={"size_bytes": len(content)},
                preview="Image file (not parsed)"
            )

    else:
        return ParsedDocument(
            content="",
            title=filename,
            file_type="unknown",
            success=False,
            error=f"No parser available for {ext}"
        )


# ============================================================================
# Utility Functions
# ============================================================================

def get_file_info(filename: str, content: bytes) -> Dict[str, Any]:
    """Get basic file info without full parsing."""
    ext = os.path.splitext(filename.lower())[1]

    return {
        "filename": filename,
        "extension": ext,
        "file_type": get_file_type(filename),
        "size_bytes": len(content),
        "size_kb": round(len(content) / 1024, 1),
        "supported": is_supported(filename)
    }


def estimate_chunks(content: str, chunk_size: int = 1000) -> int:
    """Estimate number of chunks for knowledge base."""
    if not content:
        return 0
    return max(1, len(content) // chunk_size)
